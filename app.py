"""
Multi-Site Daily Shift Report — Streamlit App
Supports: Buguruni, Puma Upanga, Puma Ocean Road, Puma Survey, India, Livingstone
Run with:  python app.py   OR   streamlit run app.py
"""

import streamlit as st
import requests
import pandas as pd
from requests.auth import HTTPDigestAuth
import urllib3
import io, sys, subprocess, os, re
from datetime import date, timedelta, datetime, timezone
from collections import defaultdict

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib import colors
from reportlab.lib.units import mm
from reportlab.platypus import (SimpleDocTemplate, Table, TableStyle,
                                Paragraph, Spacer, HRFlowable, Image as RLImage)
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.drawing.image import Image as XLImage

urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

# Tanzania uses UTC+03:00 throughout the year.
APP_TIMEZONE = timezone(timedelta(hours=3))

# Keep BAHDELA-logo.png in the same folder as app.py. The /mnt/data fallback
# is useful for local testing and does not affect Streamlit Cloud deployment.
APP_DIR = os.path.dirname(os.path.abspath(__file__))
LOGO_FILENAMES = ("BAHDELA-logo.png", "bahdela-logo.png", "BAHDELA_logo.png")

def get_logo_path():
    candidates = [os.path.join(APP_DIR, name) for name in LOGO_FILENAMES]
    candidates.extend(os.path.join("/mnt/data", name) for name in LOGO_FILENAMES)
    for candidate in candidates:
        if os.path.isfile(candidate):
            return candidate
    return None


# A later scan becomes checkout only from the allocated shift ending time.
# The grace period prevents the next workday's check-in from being used as checkout.
CHECKOUT_GRACE_HOURS = 8
SHIFT_TIME_PATTERN = re.compile(r"(\d{1,2}:\d{2})\s*-\s*(\d{1,2}:\d{2})")

def local_now():
    return datetime.now(APP_TIMEZONE)

# ── AUTO-LAUNCH ────────────────────────────────────────────────
# Only relaunch if we were NOT already started by streamlit.
# The sentinel env var _APP_LAUNCHED prevents recursive relaunching.
if os.environ.get("_APP_LAUNCHED") != "1":
    os.environ["_APP_LAUNCHED"] = "1"
    script = os.path.abspath(__file__)
    env = {**os.environ, "_APP_LAUNCHED": "1"}
    sys.exit(subprocess.call(
        ["streamlit", "run", script, "--server.headless", "false"],
        shell=(sys.platform == "win32"), env=env))

# ── PAGE CONFIG ───────────────────────────────────────────────
st.set_page_config(page_title="Shift Reports — Multi-Site", layout="wide")

# ── LIGHT-ONLY RESPONSIVE THEME ──────────────────────────────
BG      = "#FFFFFF"
SIDEBAR = "#F3EDE3"
CARD_BG = "#FFF9F2"
H1_COL  = "#5A0000"
H2_COL  = "#8B4500"
TXT     = "#2D2D2D"
MUTED   = "#6F6F6F"
INP_BG  = "#FFFFFF"
DL_BG   = "#FFFFFF"

st.markdown(f"""
<style>
  :root {{ color-scheme: light only; }}
  html, body, [class*="css"] {{ font-family:Arial,sans-serif; }}
  .stApp {{ background-color:{BG}; color:{TXT}; }}
  .block-container {{ max-width:1500px; padding-top:1.25rem; padding-bottom:2rem; }}
  [data-testid="stSidebar"] {{ background-color:{SIDEBAR}; border-right:1px solid #E3D8C8; }}

  h1,h2,h3,h4 {{ font-family:Arial,sans-serif; }}
  h1 {{ color:{H1_COL} !important; }}
  h2,h3,h4 {{ color:{H2_COL} !important; }}
  p,label,div,span {{ color:{TXT}; font-family:Arial,sans-serif; }}

  [data-testid="stMetric"] {{
    background:{CARD_BG}; border:1px solid #D8C4A7;
    border-radius:10px; padding:12px 16px;
  }}
  [data-testid="stMetricLabel"] {{ color:{H2_COL} !important; font-size:13px; }}
  [data-testid="stMetricValue"] {{ color:{H1_COL} !important; font-size:24px; font-weight:bold; }}

  /* Keep every normal button readable. */
  .stButton > button {{
    background-color:#8B0000 !important; color:#FFFFFF !important;
    border:1px solid #8B0000 !important; border-radius:7px;
    font-weight:bold; font-size:14px; padding:8px 18px;
    min-height:42px; width:100%;
  }}
  .stButton > button p,
  .stButton > button span,
  .stButton > button div {{ color:#FFFFFF !important; }}
  .stButton > button:hover {{
    background-color:#C0392B !important; border-color:#C0392B !important;
    color:#FFFFFF !important;
  }}

  /* Download buttons stay white with dark visible text. */
  .stDownloadButton > button {{
    background-color:#FFFFFF !important; color:#8B0000 !important;
    border:2px solid #8B0000 !important; border-radius:8px;
    font-size:14px; font-weight:800; padding:9px 14px;
    min-height:46px; width:100%;
    box-shadow:0 2px 6px rgba(90,0,0,.08);
  }}
  .stDownloadButton > button p,
  .stDownloadButton > button span,
  .stDownloadButton > button div {{ color:#8B0000 !important; }}
  .stDownloadButton > button:hover {{
    background-color:#FFF3CD !important; color:#5A0000 !important;
  }}

  input {{ background-color:{INP_BG} !important; color:{TXT} !important; }}
  [data-baseweb="select"] > div {{ background:#FFFFFF !important; color:{TXT} !important; }}
  hr {{ border-color:#D8C4A7 !important; }}
  .stProgress > div > div {{ background-color:#8B0000 !important; }}
  /* Make report choices look like clear, high-contrast buttons. */
  .stTabs [data-baseweb="tab-list"] {{
    display:flex !important;
    flex-wrap:wrap !important;
    gap:10px !important;
    overflow-x:visible !important;
    padding:4px 0 10px 0 !important;
  }}
  .stTabs [data-baseweb="tab"] {{
    flex:1 1 210px !important;
    justify-content:center !important;
    min-height:48px !important;
    padding:10px 16px !important;
    background:#FFFFFF !important;
    border:2px solid #8B0000 !important;
    border-radius:9px !important;
    color:#8B0000 !important;
    font-weight:800 !important;
    white-space:nowrap !important;
    box-shadow:0 2px 6px rgba(90,0,0,.08);
  }}
  .stTabs [data-baseweb="tab"] p,
  .stTabs [data-baseweb="tab"] span,
  .stTabs [data-baseweb="tab"] div {{
    color:#8B0000 !important;
    font-weight:800 !important;
  }}
  .stTabs [data-baseweb="tab"]:hover {{
    background:#FFF7F2 !important;
    border-color:#C0392B !important;
  }}
  .stTabs [data-baseweb="tab"][aria-selected="true"] {{
    background:#8B0000 !important;
    border-color:#8B0000 !important;
    color:#FFFFFF !important;
    box-shadow:0 3px 9px rgba(90,0,0,.18);
  }}
  .stTabs [data-baseweb="tab"][aria-selected="true"] p,
  .stTabs [data-baseweb="tab"][aria-selected="true"] span,
  .stTabs [data-baseweb="tab"][aria-selected="true"] div {{
    color:#FFFFFF !important;
  }}
  .stTabs [data-baseweb="tab-highlight"] {{ display:none !important; }}

  /* Site choices replace the old separate Switch buttons. */
  div[data-testid="stRadio"] [role="radiogroup"] {{
    display:flex !important; flex-wrap:wrap !important;
    gap:8px !important; align-items:stretch !important;
  }}
  div[data-testid="stRadio"] [role="radiogroup"] > label {{
    flex:1 1 145px !important; min-width:130px !important;
    margin:0 !important; padding:11px 12px !important;
    background:#FFFFFF !important; border:1px solid #D9CEC0 !important;
    border-radius:9px !important; justify-content:center !important;
    text-align:center !important; cursor:pointer !important;
  }}
  div[data-testid="stRadio"] [role="radiogroup"] > label:hover {{
    border-color:#8B0000 !important; background:#FFF9F2 !important;
  }}
  div[data-testid="stRadio"] [role="radiogroup"] > label:has(input:checked) {{
    background:#8B0000 !important; border-color:#D4A017 !important;
    box-shadow:0 2px 8px rgba(139,0,0,.14);
  }}
  div[data-testid="stRadio"] [role="radiogroup"] > label:has(input:checked) p,
  div[data-testid="stRadio"] [role="radiogroup"] > label:has(input:checked) span,
  div[data-testid="stRadio"] [role="radiogroup"] > label:has(input:checked) div {{
    color:#FFFFFF !important; font-weight:bold !important;
  }}
  div[data-testid="stRadio"] [role="radiogroup"] > label > div:first-child {{ display:none !important; }}
  div[data-testid="stRadio"] [role="radiogroup"] p {{
    width:100%; text-align:center; margin:0; font-weight:700;
  }}

  .department-grid {{
    display:grid; grid-template-columns:repeat(auto-fit,minmax(145px,1fr));
    gap:10px; margin-top:8px;
  }}
  .department-card {{
    padding:10px 8px; border-radius:7px; text-align:center;
    min-height:86px; display:flex; flex-direction:column;
    justify-content:center;
  }}

  @media (max-width: 768px) {{
    .stTabs [data-baseweb="tab"] {{
      flex:1 1 100% !important;
      min-height:50px !important;
      font-size:14px !important;
    }}
    .stDownloadButton > button {{
      min-height:48px !important;
      font-size:14px !important;
    }}
    .block-container {{ padding:0.8rem 0.7rem 1.7rem 0.7rem; }}
    h1 {{ font-size:1.65rem !important; line-height:1.2 !important; }}
    [data-testid="column"] {{ min-width:0 !important; }}
    [data-testid="stMetric"] {{ padding:10px 11px; }}
    [data-testid="stMetricValue"] {{ font-size:20px !important; }}
    div[data-testid="stRadio"] [role="radiogroup"] > label {{
      flex:1 1 calc(50% - 8px) !important; min-width:125px !important;
      padding:10px 7px !important;
    }}
    .department-grid {{ grid-template-columns:repeat(2,minmax(0,1fr)); gap:8px; }}
    .stButton > button, .stDownloadButton > button {{ min-height:46px; }}
  }}

  @media (max-width: 420px) {{
    div[data-testid="stRadio"] [role="radiogroup"] > label {{
      flex:1 1 100% !important;
    }}
    .department-grid {{ grid-template-columns:1fr; }}
  }}
</style>
""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════
# ── MULTI-SITE CONFIG ─────────────────────────────────────────
# ══════════════════════════════════════════════════════════════
SITES = {
    "Buguruni": {
        "label": "Buguruni",
        "device_ip": "217.29.138.10:4376",
        "username": "admin",
        "password": "Bahdela01",
        "dept_order": ["Shop","Manufactural","Store","Oil","Walinzi"],
        "dept_colors": {"Shop":"#8B0000","Manufactural":"#7B3F00","Store":"#B8860B",
                        "Oil":"#4B0082","Walinzi":"#555555"},
        "employees": {
            "Shop": {
                "label":"SHOP", "shift":"Shift: 08:00 - 16:30", "header_hex":"8B0000",
                "members":["Abdallah Iddi Omary","Abubakary Hussein Sharif","Ally Mohamedi Sadaka",
                           "Azizi Ibuma Issuja","Khalid Ally Bahdela","Mbarak Bates",
                           "Mohamed Omar Bahdela","Omary Bakari Athumani","Sharifu Abdallah Nyange"],
            },
            "Manufactural": {
                "label":"MANUFACTURAL", "shift":"Shift: 08:00 - 16:30", "header_hex":"7B3F00",
                "members":["Abdulkarim Juma Baraja","Hafidh Selemani Nkya","Hassani Mohamedi Kinyala",
                           "Hassani Shabani Mnadi","Hassani Swalehe Lumwe","Juma Issa Makombo",
                           "Kazeze Charles Lukindo","Salum Joshua Caleb","Shabani Bakari Mmassa",
                           "Shabani Sulemani Matola","Thadeo George Thadeo"],
            },
            "Store": {
                "label":"GODOWN", "shift":"Shift: 08:00 - 16:30", "header_hex":"B8860B",
                "members":["Ally Ally Rajabu","Edward Elinihaki Mmbara","Emily Avelini Emiliani",
                           "Ismail Masambuka Kopa","Khamis Ramadhani Masahaka","Nasibu Jafari Ngaleni",
                           "Ndensari David Ulomi","Omary Juma Chande","Said Ally Said",
                           "Said Saleh Msean","Salum Hamisi Mangochi","Shakur Ramadhani Kirungi"],
            },
            "Oil": {
                "label":"OIL", "shift":"Shifts: 06:00-14:00 | 14:00-18:00 | 18:00-06:00",
                "header_hex":"4B0082",
                "members":["Abdul Timanya Renatus","Ahmed Ally Mpogo","Azihar Mwijage Mzamiru",
                           "Bashiru Said Mgodoka","Iddi Rashidi Abdallah","Issa Sharahbil Siraji",
                           "Juma Othmani Issa","Mohamed Abdallah Hamad","Muhsin Oscar Matikila",
                           "Shali Sebe Bazo","Wilibald Antigon Urio"],
            },
            "Walinzi": {
                "label":"WALINZI (SECURITY)", "shift":"Shift: 09:00 - 09:00 (24 Hrs)",
                "header_hex":"1A1A1A",
                "members":["Anderson Mude Faru","Dickson Dickson Haule","Fadhili Ahmad Abdlaah",
                           "Fikiri Amani Abdallah","John Jofery Kapiki","Shamte Rashid Likwira"],
            },
        },
    },

    "Puma Upanga": {
        "label": "Puma Upanga",
        "device_ip": "217.29.138.128:4373",
        "username": "admin",
        "password": "Bahdela01",
        "dept_order": [
            "Oil", "Oil Supervisor", "Oil Logistics",
            "Supermarket", "Security", "Cleaner",
        ],
        "dept_colors": {
            "Oil":"#4B0082",
            "Oil Supervisor":"#7B3F00",
            "Oil Logistics":"#8B4500",
            "Supermarket":"#8B0000",
            "Security":"#1A3A5C",
            "Cleaner":"#2E6B3E",
        },
        "employees": {
            "Oil": {
                "label":"OIL", "shift":"Shifts: 07:00-14:00 | 14:00-19:00 | 19:00-06:00",
                "header_hex":"4B0082",
                "members":[
                    "Abdallah Salehe Kilindo","Abdul Jumbe Yusuf","Aisha Mustapha Losili",
                    "Asha Abdallah Ngasinda","Athumani Yusuph Sheby","Ayubu Salehe Kesi",
                    "Bakari Saidi Mchingwi","Bihasanati Saidi Nzige","Daudi Rashidi Sharabile",
                    "Elibariki Geofrey Mshana","Hamadi Adinani Iddi","Hamlati Shaibu Ghindu",
                    "Jamali Sadi Kilindo","Juma Mussa Juma","Juma Mwamedi Ngakola",
                    "Kelvin Lukemelo Mlowe","Mnyeto Miraji Rashidi",
                    "Ramadhani Abdallah Ally","Hassani Muhidini Masenga",
                    "Rashidi Majau Masai","Ramadhan Mohamed Yusuph",
                    "Rajabu Ijumaa Hizza","Nesha Habib Losiri","Msabaha Ally Msabaha",
                ],
            },
            "Oil Supervisor": {
                "label":"OIL SUPERVISOR", "shift":"Shifts: 06:00-18:00 | 18:00-06:00",
                "header_hex":"7B3F00",
                "members":[
                    "Ally Selemani Darusi","Thuweni Abdallah Hamadi","Hamisi Mzee Juma",
                ],
            },
            "Oil Logistics": {
                "label":"OIL LOGISTICS", "shift":"Shift: 08:00 - 17:00",
                "header_hex":"8B4500",
                "members":[
                    "Said Bunu Mohamed",
                ],
            },
            "Supermarket": {
                "label":"SUPERMARKET", "shift":"Shifts: 08:00-17:00 | 17:00-07:00",
                "header_hex":"8B0000",
                "members":[
                    "Avelin Said Costa","Careen Alphonce Naman","Exsadi Bonifasi Mrisho",
                    "Hadija Yasin Siraji","Happy Edwin Ndambo","Husna Ally Juma",
                    "Iddy Sufiani Msuya","James Josam Mwombeki","Latifa Rehani Abdallah",
                    "Martha Ibrahimu Mussa","Maryam Yahya Ally","Zuena Siraji Sharhabilly",
                    "Fatuma Alawi Juma",
                ],
            },
            "Security": {
                "label":"SECURITY", "shift":"Shift: 09:00 - 09:00 (24 Hrs)",
                "header_hex":"1A3A5C",
                "members":[
                    "Karim Mohamedi Radaki","Mnyeto Miraji Rashidi","Abdul Miraj Mkwizu",
                ],
            },
            "Cleaner": {
                "label":"CLEANER (HOUSE KEEPING)", "shift":"Shifts: 08:00-17:00 | 17:00-08:00",
                "header_hex":"2E6B3E",
                "members":[
                    "Ladslaus Nestor Andreas","Yunus Mohamed Zahoro",
                    "Tabaraka Mohamedi Makwati",
                ],
            },
        },
    },

    "Puma Ocean Road": {
        "label": "Puma Ocean Road",
        "device_ip": "217.29.138.127:4374",
        "username": "admin",
        "password": "Bahdela01",
        "dept_order": ["Oil","Security","Cleaner","Wakala","Accountant"],
        "dept_colors": {"Oil":"#4B0082","Security":"#1A3A5C","Cleaner":"#2E6B3E",
                        "Wakala":"#8B4500","Accountant":"#8B0000"},
        "employees": {
            "Oil": {
                "label":"OIL", "shift":"Shifts: 06:00-14:00 | 14:00-18:00 | 18:00-06:00",
                "header_hex":"4B0082",
                "members":[
                    "Abdallah Abdi Rashidi","Ally Msabaha Ally","Amina Alawi Juma",
                    "Dotto Sultan Kingalu","Esha Jamali Kessy","Juma Alawi Juma",
                    "Mahamuod Said Mgawe","Mohamed Abdallah Kambi","Mohamed Ally Mzee",
                    "Mulla Omari Kadari","Mzee Hamis Kassim","Omari Yusufu Shebly",
                    "Said Ramadhani Mnjama","Salim Siraji Bajuni","Salmin Mdhihiri Yahaya",
                    "Salmu Omari Siraji","Thomas Constans Kereza","Zahoro Rashidi Mbotoni",
                    "Emmanuel Kituki",
                ],
            },
            "Security": {
                "label":"SECURITY", "shift":"Shift: 09:00 - 09:00 (24 Hrs)",
                "header_hex":"1A3A5C",
                "members":[
                    "Mtende Hussein Jambia","Sultan Ramadhan Mikendo","Zubery Hassan Zubery",
                ],
            },
            "Cleaner": {
                "label":"CLEANER (HOUSE KEEPING)", "shift":"Shifts: 08:00-17:00 | 17:00-08:00",
                "header_hex":"2E6B3E",
                "members":[
                    "Erenest Iginasi Nyanga","Medrine Barnabas Mpangala","Samsoni Damasi Zebedayo",
                ],
            },
            "Wakala": {
                "label":"WAKALA", "shift":"Shift: 08:00 - 17:00",
                "header_hex":"8B4500",
                "members":[
                    "Nasma Twalha Mfinanga",
                ],
            },
            "Accountant": {
                "label":"ACCOUNTANT", "shift":"Shift: 08:00 - 17:00",
                "header_hex":"8B0000",
                "members":[
                    "Yusuf Abubakar Bahdela",
                ],
            },
        },
    },

    "Puma Survey": {
        "label": "Puma Survey (Savoh PUMA Sheli)",
        "device_ip": "102.205.250.241:4378",
        "username": "admin",
        "password": "Bahdela01",
        "dept_order": ["Oil","Cleaning","Supervisors","Wakala"],
        "dept_colors": {
            "Oil":"#4B0082",
            "Cleaning":"#2E6B3E",
            "Supervisors":"#7B3F00",
            "Wakala":"#8B4500",
        },
        "employees": {
            "Oil": {
                "label":"OIL", "shift":"Shifts: 07:00-14:00 | 14:00-18:00 | 19:00-06:00",
                "header_hex":"4B0082",
                "members":[
                    "Abdul Abbas Bajuni","Waziri Hatibu Waziri","Swed Khalfan Kailo",
                    "Issa Ally Juma","Zulfa Rashid Mkamba","Yassini Siraji Athumani",
                    "Yahya Mwahimu Abdul","William Joisack Maliselo","Siraji Yassini Siraji",
                    "Mustani Juberi Ngoda","Mshija Juma Lukuba","Jasmin Warid Idd",
                    "Hassani Hashimu Mpasule","Haruni Kenedy Saimoni","Haruna Muhamedi Haruna",
                    "Hamisi Abdull Siraji","France John Mapunda","Fikirini Rashidi Mussa",
                    "Baraka Raymond Daa","Abuu Jafari Salehe","Abdul Amir Mwishaha",
                    "Abdallah Ally Abdallah",
                ],
            },
            "Cleaning": {
                "label":"CLEANING / HOUSEKEEPING", "shift":"Shift: 08:00 - 17:00",
                "header_hex":"2E6B3E",
                "members":[
                    "Shabani Jafari Sobo","Saumu Omary Mwalimu","Makenya Juma Makenya",
                ],
            },
            "Supervisors": {
                "label":"SUPERVISORS / MANAGEMENT", "shift":"Shift: 08:00 - 17:00",
                "header_hex":"7B3F00",
                "members":[
                    "Omar Abdallah Bahdela","Donald Gasper Kalafya",
                ],
            },
            "Wakala": {
                "label":"WAKALA", "shift":"Shift: 08:00 - 17:00",
                "header_hex":"8B4500",
                "members":[
                    "Mwanaiddy Ally Ramadhani","Mohamed Kombo Mwinyikambi","Fatma Said Semfuko",
                ],
            },
        },
    },

    "India": {
        "label": "India",
        # Store only IP:port because fetch_events() adds the http:// prefix.
        "device_ip": "102.221.20.46:4440",
        "username": "admin",
        "password": "Bahdela01",
        "dept_order": ["Shop"],
        "dept_colors": {"Shop":"#8B0000"},
        "employees": {
            "Shop": {
                "label":"SHOP", "shift":"Shift: 08:00 - 17:00",
                "header_hex":"8B0000",
                "members":[
                    "Silivester William Yakobo",
                    "Athuman Amiri Kanda",
                    "Ally Salum Sultani",
                    "Ally Omar Mohamed",
                    "Ahmed Mohamed Bahdela",
                ],
                "member_ids": {
                    "Silivester William Yakobo":"6005",
                    "Athuman Amiri Kanda":"6004",
                    "Ally Salum Sultani":"6003",
                    "Ally Omar Mohamed":"6002",
                    "Ahmed Mohamed Bahdela":"6001",
                },
            },
        },
    },

    "Livingstone": {
        "label": "Livingstone (Outside)",
        "device_ip": "217.29.138.29:4735",
        "username": "admin",
        "password": "Bahdela01",
        "dept_order": ["Gas","Security","Cleaner"],
        "dept_colors": {"Gas":"#4B0082","Security":"#1A3A5C","Cleaner":"#2E6B3E"},
        "employees": {
            "Gas": {
                "label":"GAS", "shift":"Shift: 08:00 - 17:00",
                "header_hex":"4B0082",
                "members":[
                    "Kulwa Saidi Ndumbo","Karim Hamidu Mshana",
                    "Richard Oscar Pius","Jofrey Christian Mng'ande",
                ],
            },
            "Security": {
                "label":"SECURITY", "shift":"Shift: 08:00 - 08:00 (24 Hrs)",
                "header_hex":"1A3A5C",
                "members":[
                    "Shaibu Ahmad Chuya","Athuman Abdallah Nammenga",
                ],
            },
            "Cleaner": {
                "label":"CLEANER", "shift":"Shifts: 08:00-17:00",
                "header_hex":"2E6B3E",
                "members":[
                    "Micheal Gabriel Anthony","Nesto Atans Mdaimali",
                ],
            },
        },
    },
}

# ── HARD CONFIG CLEANUP ──────────────────────────────────────
def sanitize_puma_upanga_config():
    """Keep Upanga Manager-free and ensure Maryam appears once in Supermarket."""
    upanga = SITES.get("Puma Upanga", {})
    employees = upanga.get("employees", {})

    # Remove the retired Manager department from every configuration source.
    employees.pop("Manager", None)
    upanga["dept_order"] = [
        dept for dept in upanga.get("dept_order", []) if dept != "Manager"
    ]
    upanga.get("dept_colors", {}).pop("Manager", None)

    # Remove Maryam from all Upanga groups first, then add her once to Supermarket.
    maryam = "Maryam Yahya Ally"
    for department in employees.values():
        department["members"] = [
            name for name in department.get("members", [])
            if name.strip().casefold() != maryam.casefold()
        ]

    supermarket = employees.get("Supermarket")
    if supermarket is not None:
        supermarket.setdefault("members", []).append(maryam)


sanitize_puma_upanga_config()

# ── Runtime site selection (resolved after sidebar renders) ───
# These are set dynamically in the UI section below; defaults here
ALL_EMPLOYEES = SITES["Buguruni"]["employees"]
DEPT_ORDER    = SITES["Buguruni"]["dept_order"]
DEPT_COLORS   = SITES["Buguruni"]["dept_colors"]

# ── CORE HELPERS ──────────────────────────────────────────────
def get_dept(name):
    nl=name.lower().strip(); np=name.strip().split()
    for dk,dept in ALL_EMPLOYEES.items():
        for m in dept["members"]:
            if m.lower()==nl: return dk
            mp=m.split()
            if len(np)>=2 and len(mp)>=2 and np[0]==mp[0] and np[1]==mp[1]: return dk
    return None

def get_configured_employee_id(dept_key, name):
    """Return an employee ID configured for an absent/static member, if available."""
    department = ALL_EMPLOYEES.get(dept_key, {})
    member_ids = department.get("member_ids", {})
    wanted = str(name).strip().casefold()
    for member_name, employee_id in member_ids.items():
        if str(member_name).strip().casefold() == wanted:
            return str(employee_id)
    return "-"

def to_min(t):
    if not t or t=="-": return None
    try: h,m=t.split(":"); return int(h)*60+int(m)
    except: return None

def get_shift_pairs(dept_key):
    """Read all configured start/end time pairs for a department."""
    shift_text = str(ALL_EMPLOYEES.get(dept_key, {}).get("shift", ""))
    pairs = []
    for start_text, end_text in SHIFT_TIME_PATTERN.findall(shift_text):
        start_minute = to_min(start_text)
        end_minute = to_min(end_text)
        if start_minute is not None and end_minute is not None:
            pairs.append((start_minute, end_minute))
    return pairs


def get_shift_window(dept_key, check_in_dt):
    """Return the allocated shift start and end nearest to the check-in scan."""
    pairs = get_shift_pairs(dept_key) or [(8 * 60, 17 * 60)]
    check_in_minute = check_in_dt.hour * 60 + check_in_dt.minute

    def circular_distance(start_minute):
        difference = abs(check_in_minute - start_minute)
        return min(difference, 1440 - difference)

    start_minute, end_minute = min(
        pairs,
        key=lambda pair: circular_distance(pair[0]),
    )

    signed_offset = ((start_minute - check_in_minute + 720) % 1440) - 720
    midnight = check_in_dt.replace(hour=0, minute=0, second=0, microsecond=0)
    shift_start = midnight + timedelta(minutes=check_in_minute + signed_offset)

    duration_minutes = (end_minute - start_minute) % 1440
    if duration_minutes == 0:
        duration_minutes = 1440
    shift_end = shift_start + timedelta(minutes=duration_minutes)
    return shift_start, shift_end


def parse_event_datetime(value):
    """Parse device timestamps and normalize them to Tanzania time."""
    raw = str(value or "").strip()
    if not raw:
        return None
    try:
        parsed = datetime.fromisoformat(raw.replace("Z", "+00:00"))
    except ValueError:
        try:
            parsed = datetime.strptime(raw[:19], "%Y-%m-%dT%H:%M:%S")
        except ValueError:
            return None
    if parsed.tzinfo is None:
        parsed = parsed.replace(tzinfo=APP_TIMEZONE)
    return parsed.astimezone(APP_TIMEZONE)


def checkin_fill(dept_key,check_in):
    ci=to_min(check_in)
    if ci is None: return "EEEEEE","999999"
    # Oil shifts: 06:00(360) or 07:00(420) start, 14:00(840), 19:00(1140)
    if dept_key in ("Oil",):
        if ci < 840:   ss = 360   # morning shift start 06:00 or 07:00
        elif ci < 1140: ss = 840  # afternoon 14:00
        else:           ss = 1140 # night 19:00
        if ci < ss:      return "C6EFCE","276221"
        if ci > ss+60:   return "FFCCCC","CC0000"
        return "FFFFFF","333333"
    # Oil Supervisor: 06:00-18:00 or 18:00-06:00
    if dept_key == "Oil Supervisor":
        ss = 360 if ci < 720 else 1080
        if ci < ss:      return "C6EFCE","276221"
        if ci > ss+60:   return "FFCCCC","CC0000"
        return "FFFFFF","333333"
    # Oil Logistics: 08:00-17:00
    if dept_key == "Oil Logistics":
        ss = 480
        if ci < ss:      return "C6EFCE","276221"
        if ci > ss+60:   return "FFCCCC","CC0000"
        return "FFFFFF","333333"
    # Security/Walinzi: 24hr or two-shift
    if dept_key in ("Walinzi","Security"):
        ss = 480 if 480 <= ci < 960 else 960
        if ci < ss:     return "C6EFCE","276221"
        if ci > ss+60:  return "FFCCCC","CC0000"
        return "FFFFFF","333333"
    # Supermarket: 08:00-17:00 or 17:00-07:00
    if dept_key == "Supermarket":
        ss = 480 if ci < 1020 else 1020
        if ci < ss:     return "C6EFCE","276221"
        if ci > ss+60:  return "FFCCCC","CC0000"
        return "FFFFFF","333333"
    # Default day-shift depts (Shop, Manufactural, Store, Cleaner, Wakala, Manager, Accountant)
    if ci < 495: return "C6EFCE","276221"
    if ci > 540: return "FFCCCC","CC0000"
    return "FFFFFF","333333"

def style_viewer_dataframe(dataframe, dept_key):
    """Apply the report's attendance colours to the on-screen viewer."""
    columns = list(dataframe.columns)
    check_in_index = columns.index("Check In")

    def style_row(row):
        styles = ["" for _ in columns]

        # Absent staff use the same grey styling as exported reports.
        if row.get("Status") == "Absent":
            return [
                "background-color: #EEEEEE; color: #999999;"
                for _ in columns
            ]

        check_in = str(row.get("Check In") or "").strip()
        if check_in:
            background, text_colour = checkin_fill(dept_key, check_in)
            styles[check_in_index] = (
                f"background-color: #{background}; "
                f"color: #{text_colour}; font-weight: 700;"
            )
        return styles

    return (
        dataframe.style
        .apply(style_row, axis=1)
        .set_properties(**{
            "border-color": "#D4A017",
            "vertical-align": "middle",
        })
    )

def hex2rgb(h):
    h=h.lstrip("#"); return tuple(int(h[i:i+2],16) for i in (0,2,4))
def hex2rl(h):
    r,g,b=hex2rgb(h); return colors.Color(r/255,g/255,b/255)

def calc_mins(ci,co):
    try:
        if co=="-": return 0
        t1=datetime.strptime(ci,"%H:%M"); t2=datetime.strptime(co,"%H:%M")
        if t2<t1: t2+=timedelta(days=1)
        return int((t2-t1).total_seconds())//60
    except: return 0

def calc_hours(ci,co):
    m=calc_mins(ci,co)
    if m==0 and co=="-": return "-"
    h,mm=divmod(m,60); return f"{h} hrs {mm:02d} min"

def parse_by_day(events, report_start=None, report_end=None):
    """
    Build attendance rows from chronological scans.

    Checkout is accepted only at the allocated shift ending time or later.
    Repeated scans before shift end are ignored, so checkout and worked hours
    remain blank when there is no valid checkout scan.
    """
    employee_scans = defaultdict(list)
    employee_names = {}

    for event in events:
        employee_id = str(
            event.get("employeeNoString")
            or event.get("employeeNo")
            or event.get("cardNo")
            or ""
        ).strip()
        name = str(event.get("name") or "").strip()
        if not employee_id and not name:
            continue
        if not employee_id:
            employee_id = name

        scan_time = parse_event_datetime(event.get("time"))
        if scan_time is None:
            continue

        employee_names[employee_id] = name or f"ID {employee_id}"
        employee_scans[employee_id].append(scan_time)

    rows = []
    grace = timedelta(hours=CHECKOUT_GRACE_HOURS)

    for employee_id, scans in employee_scans.items():
        scans = sorted(set(scans))
        employee_name = employee_names[employee_id]
        department_key = get_dept(employee_name)
        index = 0

        while index < len(scans):
            check_in_dt = scans[index]
            _, allocated_end = get_shift_window(department_key, check_in_dt)

            checkout_dt = None
            checkout_index = None
            next_index = index + 1
            candidate_index = index + 1

            while candidate_index < len(scans):
                candidate = scans[candidate_index]

                # Ignore duplicate/extra scans made before the shift ends.
                if candidate < allocated_end:
                    next_index = candidate_index + 1
                    candidate_index += 1
                    continue

                # The first scan at shift end or later is checkout, provided it
                # is still reasonably close to that allocated shift.
                if candidate <= allocated_end + grace:
                    checkout_dt = candidate
                    checkout_index = candidate_index
                break

            report_day = check_in_dt.date()
            within_start = report_start is None or report_day >= report_start
            within_end = report_end is None or report_day <= report_end

            if within_start and within_end:
                check_in = check_in_dt.strftime("%H:%M")
                check_out = checkout_dt.strftime("%H:%M") if checkout_dt else "-"
                rows.append({
                    "name": employee_name,
                    "employee_id": employee_id,
                    "date": report_day.isoformat(),
                    "check_in": check_in,
                    "check_out": check_out,
                    "hours": calc_hours(check_in, check_out),
                    "mins": calc_mins(check_in, check_out),
                })

            index = (
                checkout_index + 1
                if checkout_index is not None
                else max(next_index, index + 1)
            )

    # Keep only one daily record per employee and prefer a record with checkout.
    consolidated = {}
    for row in sorted(rows, key=lambda item: (item["date"], item["name"], item["check_in"])):
        key = (row["employee_id"], row["date"])
        if key not in consolidated:
            consolidated[key] = row
            continue
        existing = consolidated[key]
        if existing["check_out"] == "-" and row["check_out"] != "-":
            existing["check_out"] = row["check_out"]
            existing["hours"] = calc_hours(existing["check_in"], existing["check_out"])
            existing["mins"] = calc_mins(existing["check_in"], existing["check_out"])

    return sorted(consolidated.values(), key=lambda row: (row["date"], row["name"]))

def build_dept_data(day_rows):
    result={}
    for dk in DEPT_ORDER:
        present=sorted([r for r in day_rows if get_dept(r["name"])==dk],
                       key=lambda r: to_min(r["check_in"]) or 9999)
        pl=set()
        for r in present:
            pl.add(r["name"].lower().strip()); p=r["name"].strip().split()
            if len(p)>=2: pl.add(f"{p[0]} {p[1]}".lower())
        absent=[]
        for mname in ALL_EMPLOYEES[dk]["members"]:
            ml=mname.lower().strip(); mp=mname.strip().split()
            fl=f"{mp[0]} {mp[1]}".lower() if len(mp)>=2 else ml
            if ml not in pl and fl not in pl: absent.append(mname)
        result[dk]={"present":present,"absent":absent}
    return result

# Aggregate multi-day: per employee first/last per day, then summarise
def build_summary(rows_by_date, start_date, end_date):
    """Returns per-dept summary: {dk: [{name, id, days_present, avg_in, avg_out, total_hrs, days_absent}]}"""
    all_days=[]
    d=start_date
    while d<=end_date: all_days.append(d.strftime("%Y-%m-%d")); d+=timedelta(days=1)
    total_days=len(all_days)

    # collect per employee
    emp_data=defaultdict(lambda:{"name":"","days":{}})
    for day_str,day_rows in rows_by_date.items():
        for r in day_rows:
            eid=r["employee_id"]
            emp_data[eid]["name"]=r["name"]
            emp_data[eid]["days"][day_str]=r

    result={}
    for dk in DEPT_ORDER:
        members=ALL_EMPLOYEES[dk]["members"]
        dept_rows=[]
        seen_ids=set()
        for mname in members:
            # find matching emp_id
            eid=None; edata=None
            nl=mname.lower().strip(); mp=mname.strip().split()
            for eid_,edata_ in emp_data.items():
                n=edata_["name"].lower().strip(); p=edata_["name"].strip().split()
                if n==nl or (len(mp)>=2 and len(p)>=2 and mp[0]==p[0] and mp[1]==p[1]):
                    eid=eid_; edata=edata_; break
            if eid and edata:
                seen_ids.add(eid)
                days_rec=edata["days"]
                days_present=len(days_rec)
                ins=[r["check_in"] for r in days_rec.values() if r["check_in"]!="-"]
                outs=[r["check_out"] for r in days_rec.values() if r["check_out"]!="-"]
                total_mins=sum(r["mins"] for r in days_rec.values())
                h,m=divmod(total_mins,60)
                dept_rows.append({
                    "name":edata["name"],"id":eid,
                    "days_present":days_present,"days_absent":total_days-days_present,
                    "avg_in":avg_time(ins),"avg_out":avg_time(outs),
                    "total_hrs":f"{h}h {m:02d}m" if total_mins>0 else "-",
                    "total_mins":total_mins,
                })
            else:
                dept_rows.append({
                    "name":mname,"id":get_configured_employee_id(dk, mname),
                    "days_present":0,"days_absent":total_days,
                    "avg_in":"-","avg_out":"-","total_hrs":"-","total_mins":0,
                })
        dept_rows.sort(key=lambda r:(-r["days_present"],r["name"]))
        result[dk]=dept_rows
    return result,total_days

def avg_time(tl):
    try:
        mins=[int(t[:2])*60+int(t[3:5]) for t in tl if t and t!="-"]
        if not mins: return "-"
        a=sum(mins)//len(mins); return f"{a//60:02d}:{a%60:02d}"
    except: return "-"

# ── FETCH ─────────────────────────────────────────────────────
def fetch_events(device_ip,username,password,start_date,end_date,progress_cb=None):
    url=f"http://{device_ip}/ISAPI/AccessControl/AcsEvent?format=json"
    auth=HTTPDigestAuth(username,password)
    all_events,position,page=[],0,1
    while True:
        payload={"AcsEventCond":{"searchID":"1","searchResultPosition":position,
            "maxResults":30,"major":0,"minor":0,
            "startTime":f"{start_date}T00:00:00+03:00",
            "endTime":f"{end_date}T23:59:59+03:00"}}
        try: r=requests.post(url,auth=auth,json=payload,timeout=15,verify=False)
        except requests.exceptions.RequestException as e: return None,str(e)
        if r.status_code!=200: return None,f"HTTP {r.status_code}: {r.text[:200]}"
        block=r.json().get("AcsEvent",{}); events=block.get("InfoList",[])
        total=block.get("totalMatches",0); num=block.get("numOfMatches",0)
        if not events: break
        all_events.extend(events); position+=num
        if progress_cb and total>0: progress_cb(min(position/total,0.55),f"Fetching... {position}/{total} events")
        if position>=total or num<30: break
        page+=1
    attendance=[e for e in all_events if e.get("employeeNoString") or e.get("name")]
    return attendance,None

# ══════════════════════════════════════════════════════════════
# ── DOCX BUILDER ─────────────────────────────────────────────
# ══════════════════════════════════════════════════════════════
def _bg(cell,h):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")): tcPr.remove(old)
    shd=OxmlElement("w:shd"); shd.set(qn("w:val"),"clear")
    shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),h.upper()); tcPr.append(shd)

def _bdr(cell,color="D4A017",sz="4"):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcBorders")): tcPr.remove(old)
    b=OxmlElement("w:tcBorders")
    for s in ("top","left","bottom","right"):
        el=OxmlElement(f"w:{s}"); el.set(qn("w:val"),"single")
        el.set(qn("w:sz"),sz); el.set(qn("w:color"),color); b.append(el)
    tcPr.append(b)

def _mar(cell):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcMar")): tcPr.remove(old)
    mar=OxmlElement("w:tcMar")
    for s,v in (("top",60),("left",120),("bottom",60),("right",120)):
        el=OxmlElement(f"w:{s}"); el.set(qn("w:w"),str(v)); el.set(qn("w:type"),"dxa"); mar.append(el)
    tcPr.append(mar)

def _wc(cell,text,bold=False,size=9,color="2D2D2D",align=WD_ALIGN_PARAGRAPH.LEFT,bg=None,border="D4A017"):
    cell.vertical_alignment=WD_ALIGN_VERTICAL.CENTER; _bdr(cell,color=border); _mar(cell)
    if bg: _bg(cell,bg)
    p=cell.paragraphs[0]; p.alignment=align
    p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0)
    run=p.add_run(str(text)); run.bold=bold; run.font.name="Arial"
    run.font.size=Pt(size); run.font.color.rgb=RGBColor(*hex2rgb(color))

def _rh(row,cm=0.6):
    tr=row._tr; trPr=tr.get_or_add_trPr()
    trH=OxmlElement("w:trHeight"); trH.set(qn("w:val"),str(int(cm*567)))
    trH.set(qn("w:hRule"),"atLeast"); trPr.append(trH)

def _cw(table,col,cm):
    for row in table.rows: row.cells[col].width=Cm(cm)

def _dx_banner(doc,label,shift,bg):
    p=doc.add_paragraph()
    p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(0)
    pPr=p._p.get_or_add_pPr()
    shd=OxmlElement("w:shd"); shd.set(qn("w:val"),"clear")
    shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),bg.upper()); pPr.append(shd)
    r1=p.add_run(f"  {label}"); r1.bold=True; r1.font.name="Arial"
    r1.font.size=Pt(13); r1.font.color.rgb=RGBColor(*hex2rgb("FFF3CD"))
    r2=p.add_run(f"   -   {shift}"); r2.font.name="Arial"
    r2.font.size=Pt(9); r2.font.color.rgb=RGBColor(*hex2rgb("D4A017"))

def _docx_title(doc,title_text,subtitle_text):
    doc.styles["Normal"].paragraph_format.space_after=Pt(0)
    logo_path = get_logo_path()
    if logo_path:
        logo_paragraph = doc.add_paragraph()
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_paragraph.paragraph_format.space_before = Pt(0)
        logo_paragraph.paragraph_format.space_after = Pt(3)
        logo_run = logo_paragraph.add_run()
        logo_run.add_picture(logo_path, width=Cm(4.2))
    doc.styles["Normal"].paragraph_format.space_before=Pt(0)
    tp=doc.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER; tp.paragraph_format.space_after=Pt(4)
    r=tp.add_run(title_text); r.bold=True; r.font.name="Arial"
    r.font.size=Pt(20); r.font.color.rgb=RGBColor(*hex2rgb("8B0000"))
    sp=doc.add_paragraph(); sp.alignment=WD_ALIGN_PARAGRAPH.CENTER; sp.paragraph_format.space_after=Pt(4)
    r=sp.add_run(subtitle_text); r.font.name="Arial"
    r.font.size=Pt(12); r.font.color.rgb=RGBColor(*hex2rgb("B8860B"))
    hr=doc.add_paragraph(); hr.paragraph_format.space_after=Pt(8)
    pPr=hr._p.get_or_add_pPr(); pBdr=OxmlElement("w:pBdr")
    bot=OxmlElement("w:bottom"); bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),"12")
    bot.set(qn("w:color"),"8B0000"); pBdr.append(bot); pPr.append(pBdr)

def build_docx_daily(rows_by_date,start_date,end_date,site="Buguruni"):
    doc=Document()
    sec=doc.sections[0]; sec.page_width=Cm(29.7); sec.page_height=Cm(21.0)
    sec.left_margin=sec.right_margin=Cm(1.5); sec.top_margin=sec.bottom_margin=Cm(1.5)
    dlabel=(start_date.strftime("%d %B %Y") if start_date==end_date
            else f"{start_date.strftime('%d %b')} - {end_date.strftime('%d %b %Y')}")
    _docx_title(doc,"DAILY SHIFT REPORT",f"Site: {site}   -   {dlabel}")
    lp=doc.add_paragraph(); lp.paragraph_format.space_after=Pt(10)
    for text,col,bold in [("Check-In Key:   ","444444",True),("  Early  ","276221",False),
                           ("  On Time  ","555555",False),("  Late  ","CC0000",False),("  Absent","999999",False)]:
        r=lp.add_run(text); r.bold=bold; r.font.name="Arial"
        r.font.size=Pt(9); r.font.color.rgb=RGBColor(*hex2rgb(col))
    total=0; day=start_date
    DX=[1.1,8.2,1.8,2.3,2.3,3.5]
    HDR=["#","Employee Name","ID","Check In","Check Out","Hours Worked"]
    DAL=[WD_ALIGN_PARAGRAPH.CENTER,WD_ALIGN_PARAGRAPH.LEFT,WD_ALIGN_PARAGRAPH.CENTER,
         WD_ALIGN_PARAGRAPH.CENTER,WD_ALIGN_PARAGRAPH.CENTER,WD_ALIGN_PARAGRAPH.CENTER]
    while day<=end_date:
        if (end_date-start_date).days>0:
            dp=doc.add_paragraph(); dp.paragraph_format.space_before=Pt(8); dp.paragraph_format.space_after=Pt(2)
            r=dp.add_run(day.strftime("%A, %d %B %Y")); r.bold=True; r.font.name="Arial"
            r.font.size=Pt(11); r.font.color.rgb=RGBColor(*hex2rgb("8B0000"))
        day_rows=rows_by_date.get(day.strftime("%Y-%m-%d"),[])
        total+=len(day_rows); dept_data=build_dept_data(day_rows)
        for dk in DEPT_ORDER:
            d=ALL_EMPLOYEES[dk]; pd=dept_data[dk]
            if not pd["present"] and not pd["absent"]: continue
            _dx_banner(doc,d["label"],d["shift"],d["header_hex"])
            tbl=doc.add_table(rows=1,cols=6); tbl.style="Table Grid"
            for i,w in enumerate(DX): _cw(tbl,i,w)
            row=tbl.rows[0]; _rh(row,0.75)
            for i,(l,a) in enumerate(zip(HDR,DAL)):
                _wc(row.cells[i],l,bold=True,size=9.5,color="FFF3CD",align=a,bg="C0392B",border="8B0000")
            for idx,rec in enumerate(pd["present"]):
                rw=tbl.add_row(); _rh(rw,0.6)
                fill="FFFFFF"
                ci_bg,ci_tc=checkin_fill(dk,rec["check_in"])
                _wc(rw.cells[0],idx+1,align=WD_ALIGN_PARAGRAPH.CENTER,bg=fill)
                _wc(rw.cells[1],rec["name"],bg=fill)
                _wc(rw.cells[2],rec["employee_id"],align=WD_ALIGN_PARAGRAPH.CENTER,bg=fill)
                _wc(rw.cells[3],rec["check_in"],bold=True,color=ci_tc,align=WD_ALIGN_PARAGRAPH.CENTER,bg=ci_bg)
                _wc(rw.cells[4],rec["check_out"] if rec["check_out"]!="-" else "",align=WD_ALIGN_PARAGRAPH.CENTER,bg=fill)
                _wc(rw.cells[5],rec["hours"] if rec["hours"]!="-" else "",align=WD_ALIGN_PARAGRAPH.CENTER,bg=fill)
            for j,nm in enumerate(pd["absent"]):
                rw=tbl.add_row(); _rh(rw,0.6)
                configured_id = get_configured_employee_id(dk, nm)
                for i,(t,a) in enumerate(zip([len(pd["present"])+j+1,nm,configured_id,"-","-","-"],DAL)):
                    _wc(rw.cells[i],t,color="999999",align=a,bg="EEEEEE")
            doc.add_paragraph().paragraph_format.space_after=Pt(4)
        day+=timedelta(days=1)
    fp=doc.add_paragraph(); fp.paragraph_format.space_before=Pt(10)
    pPr=fp._p.get_or_add_pPr(); pBdr=OxmlElement("w:pBdr")
    top=OxmlElement("w:top"); top.set(qn("w:val"),"single"); top.set(qn("w:sz"),"10")
    top.set(qn("w:color"),"8B0000"); pBdr.append(top); pPr.append(pBdr)
    r=fp.add_run(f"Total records: {total}   "); r.bold=True
    r.font.name="Arial"; r.font.size=Pt(10); r.font.color.rgb=RGBColor(*hex2rgb("8B0000"))
    r=fp.add_run(f"Report generated: {datetime.today().strftime('%d %B %Y')}")
    r.font.name="Arial"; r.font.size=Pt(9); r.font.color.rgb=RGBColor(*hex2rgb("B8860B"))
    buf=io.BytesIO(); doc.save(buf); buf.seek(0); return buf.read()

def build_docx_summary(summary,total_days,start_date,end_date,label,site="Buguruni"):
    doc=Document()
    sec=doc.sections[0]; sec.page_width=Cm(29.7); sec.page_height=Cm(21.0)
    sec.left_margin=sec.right_margin=Cm(1.5); sec.top_margin=sec.bottom_margin=Cm(1.5)
    dlabel=f"{start_date.strftime('%d %b')} - {end_date.strftime('%d %b %Y')}  ({total_days} days)"
    _docx_title(doc,f"{label.upper()} SUMMARY REPORT",f"Site: {site}   -   {dlabel}")
    DX=[1.0,7.5,1.8,2.0,2.0,2.0,2.5,2.0]
    HDR=["#","Employee Name","ID","Days Present","Days Absent","Avg Check-In","Avg Check-Out","Total Hours"]
    DAL=[WD_ALIGN_PARAGRAPH.CENTER,WD_ALIGN_PARAGRAPH.LEFT,WD_ALIGN_PARAGRAPH.CENTER,
         WD_ALIGN_PARAGRAPH.CENTER,WD_ALIGN_PARAGRAPH.CENTER,WD_ALIGN_PARAGRAPH.CENTER,
         WD_ALIGN_PARAGRAPH.CENTER,WD_ALIGN_PARAGRAPH.CENTER]
    for dk in DEPT_ORDER:
        d=ALL_EMPLOYEES[dk]; rows=summary[dk]
        if not rows: continue
        _dx_banner(doc,d["label"],d["shift"],d["header_hex"])
        tbl=doc.add_table(rows=1,cols=8); tbl.style="Table Grid"
        for i,w in enumerate(DX): _cw(tbl,i,w)
        hrow=tbl.rows[0]; _rh(hrow,0.75)
        for i,(l,a) in enumerate(zip(HDR,DAL)):
            _wc(hrow.cells[i],l,bold=True,size=9,color="FFF3CD",align=a,bg="C0392B",border="8B0000")
        for idx,rec in enumerate(rows):
            rw=tbl.add_row(); _rh(rw,0.6)
            fill="FFFFFF"
            ab_col="CC0000" if rec["days_absent"]>0 else "2D2D2D"
            if rec["days_present"]==0: fill="EEEEEE"
            _wc(rw.cells[0],idx+1,align=WD_ALIGN_PARAGRAPH.CENTER,bg=fill)
            _wc(rw.cells[1],rec["name"],bg=fill,color="999999" if rec["days_present"]==0 else "2D2D2D")
            _wc(rw.cells[2],rec["id"],align=WD_ALIGN_PARAGRAPH.CENTER,bg=fill,color="999999" if rec["days_present"]==0 else "2D2D2D")
            _wc(rw.cells[3],str(rec["days_present"]),bold=True,align=WD_ALIGN_PARAGRAPH.CENTER,bg=fill,color="276221" if rec["days_present"]>0 else "999999")
            _wc(rw.cells[4],str(rec["days_absent"]),bold=(rec["days_absent"]>0),align=WD_ALIGN_PARAGRAPH.CENTER,bg=fill,color=ab_col)
            _wc(rw.cells[5],rec["avg_in"],align=WD_ALIGN_PARAGRAPH.CENTER,bg=fill)
            _wc(rw.cells[6],rec["avg_out"],align=WD_ALIGN_PARAGRAPH.CENTER,bg=fill)
            _wc(rw.cells[7],rec["total_hrs"],align=WD_ALIGN_PARAGRAPH.CENTER,bg=fill)
        doc.add_paragraph().paragraph_format.space_after=Pt(4)
    buf=io.BytesIO(); doc.save(buf); buf.seek(0); return buf.read()

# ══════════════════════════════════════════════════════════════
# ── PDF BUILDER ──────────────────────────────────────────────
# ══════════════════════════════════════════════════════════════
def RP(text,bold=False,size=8,color=None,align=TA_LEFT):
    if color is None: color=colors.black
    return Paragraph(str(text),ParagraphStyle("x",
        fontName="Helvetica-Bold" if bold else "Helvetica",
        fontSize=size,textColor=color,alignment=align,leading=size+3))

def _pdf_doc(buf):
    PW,PH=landscape(A4); LM=RM=15*mm; TM=BM=12*mm
    return SimpleDocTemplate(buf,pagesize=landscape(A4),
        leftMargin=LM,rightMargin=RM,topMargin=TM,bottomMargin=BM), PW-LM-RM

def _pdf_banner(d,W):
    bg=hex2rl(d["header_hex"])
    t=Table([[RP(f"  {d['label']}",bold=True,size=11,color=hex2rl("FFF3CD")),
              RP(d["shift"],size=8,color=hex2rl("D4A017"))]],colWidths=[W*0.28,W*0.72])
    t.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,-1),bg),("VALIGN",(0,0),(-1,-1),"MIDDLE"),
        ("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5),
        ("LEFTPADDING",(0,0),(-1,-1),4),("RIGHTPADDING",(0,0),(-1,-1),4)]))
    return t

def _pdf_table(data,style_cmds,cw):
    t=Table(data,colWidths=cw); t.setStyle(TableStyle(style_cmds)); return t

def _pdf_logo():
    logo_path = get_logo_path()
    if not logo_path:
        return None
    logo = RLImage(logo_path)
    target_width = 34 * mm
    aspect_ratio = logo.imageHeight / float(logo.imageWidth or 1)
    logo.drawWidth = target_width
    logo.drawHeight = target_width * aspect_ratio
    logo.hAlign = "CENTER"
    return logo

def build_pdf_daily(rows_by_date,start_date,end_date,site="Buguruni"):
    buf=io.BytesIO(); doc,W=_pdf_doc(buf); story=[]
    CR=hex2rl("8B0000"); CRM=hex2rl("C0392B"); CYL=hex2rl("B8860B")
    CYLL=hex2rl("FFF3CD"); CCR=hex2rl("FFFFFF"); CGB=hex2rl("EEEEEE")
    CGT=hex2rl("999999"); CMG=hex2rl("CCCCCC")
    dlabel=(start_date.strftime("%d %B %Y") if start_date==end_date
            else f"{start_date.strftime('%d %b')} - {end_date.strftime('%d %b %Y')}")
    logo = _pdf_logo()
    if logo:
        story.append(logo)
        story.append(Spacer(1,2))
    story.append(RP("DAILY SHIFT REPORT",bold=True,size=16,color=CR,align=TA_CENTER))
    story.append(Spacer(1,4))
    story.append(RP(f"Site: {site}   -   {dlabel}",size=10,color=CYL,align=TA_CENTER))
    story.append(Spacer(1,4))
    story.append(HRFlowable(width=W,thickness=1.5,color=CR,spaceAfter=6))
    leg=[[RP("  Early",size=8,color=hex2rl("276221")),RP("  On Time",size=8),
          RP("  Late",size=8,color=hex2rl("CC0000")),RP("  Absent",size=8,color=CGT)]]
    lt=Table(leg,colWidths=[W*0.12,W*0.14,W*0.1,W*0.64])
    lt.setStyle(TableStyle([("VALIGN",(0,0),(-1,-1),"MIDDLE"),("BOTTOMPADDING",(0,0),(-1,-1),4)]))
    story.append(lt)
    cw=[W*0.04,W*0.31,W*0.08,W*0.12,W*0.12,W*0.15]
    total=0; day=start_date
    while day<=end_date:
        if (end_date-start_date).days>0:
            story.append(Spacer(1,8))
            story.append(RP(day.strftime("%A, %d %B %Y"),bold=True,size=10,color=CR))
        day_rows=rows_by_date.get(day.strftime("%Y-%m-%d"),[])
        total+=len(day_rows); dept_data=build_dept_data(day_rows)
        for dk in DEPT_ORDER:
            d=ALL_EMPLOYEES[dk]; pd=dept_data[dk]
            if not pd["present"] and not pd["absent"]: continue
            story.append(Spacer(1,4))
            story.append(_pdf_banner(d,W))
            data=[[RP("#",bold=True,size=8,color=CYLL,align=TA_CENTER),
                   RP("Employee Name",bold=True,size=8,color=CYLL),
                   RP("ID",bold=True,size=8,color=CYLL,align=TA_CENTER),
                   RP("Check In",bold=True,size=8,color=CYLL,align=TA_CENTER),
                   RP("Check Out",bold=True,size=8,color=CYLL,align=TA_CENTER),
                   RP("Hours Worked",bold=True,size=8,color=CYLL,align=TA_CENTER)]]
            style=[("BACKGROUND",(0,0),(-1,0),CRM),("BOX",(0,0),(-1,-1),0.5,CMG),
                   ("INNERGRID",(0,0),(-1,-1),0.5,CMG),("VALIGN",(0,0),(-1,-1),"MIDDLE"),
                   ("TOPPADDING",(0,0),(-1,-1),4),("BOTTOMPADDING",(0,0),(-1,-1),4),
                   ("LEFTPADDING",(0,0),(-1,-1),5),("RIGHTPADDING",(0,0),(-1,-1),5)]
            for idx,rec in enumerate(pd["present"]):
                fill=CCR
                ci_bg_h,ci_tc_h=checkin_fill(dk,rec["check_in"])
                ci_bg=hex2rl(ci_bg_h); ci_tc=hex2rl(ci_tc_h)
                data.append([RP(str(idx+1),size=8,align=TA_CENTER),RP(rec["name"],size=8),
                              RP(rec["employee_id"],size=8,align=TA_CENTER),
                              RP(rec["check_in"],bold=True,size=8,color=ci_tc,align=TA_CENTER),
                              RP(rec["check_out"] if rec["check_out"]!="-" else "",size=8,align=TA_CENTER),
                              RP(rec["hours"] if rec["hours"]!="-" else "",size=8,align=TA_CENTER)])
                ri=len(data)-1
                style.append(("BACKGROUND",(0,ri),(-1,ri),fill))
                style.append(("BACKGROUND",(3,ri),(3,ri),ci_bg))
            for j,nm in enumerate(pd["absent"]):
                configured_id = get_configured_employee_id(dk, nm)
                data.append([RP(str(len(pd["present"])+j+1),size=8,color=CGT,align=TA_CENTER),
                              RP(nm,size=8,color=CGT),RP(configured_id,size=8,color=CGT,align=TA_CENTER),
                              RP("-",size=8,color=CGT,align=TA_CENTER),RP("-",size=8,color=CGT,align=TA_CENTER),
                              RP("-",size=8,color=CGT,align=TA_CENTER)])
                ri=len(data)-1
                style.append(("BACKGROUND",(0,ri),(-1,ri),CGB))
            story.append(_pdf_table(data,style,cw))
        day+=timedelta(days=1)
    story.append(Spacer(1,8))
    story.append(HRFlowable(width=W,thickness=1,color=CR,spaceAfter=4))
    story.append(RP(f"Total records: {total}   |   Generated: {datetime.today().strftime('%d %B %Y')}",size=8,color=CYL))
    doc.build(story); buf.seek(0); return buf.read()

def build_pdf_summary(summary,total_days,start_date,end_date,label,site="Buguruni"):
    buf=io.BytesIO(); doc,W=_pdf_doc(buf); story=[]
    CR=hex2rl("8B0000"); CRM=hex2rl("C0392B"); CYL=hex2rl("B8860B")
    CYLL=hex2rl("FFF3CD"); CCR=hex2rl("FFFFFF"); CGB=hex2rl("EEEEEE")
    CGT=hex2rl("999999"); CMG=hex2rl("CCCCCC")
    CGN=hex2rl("276221"); CRD=hex2rl("CC0000")
    dlabel=f"{start_date.strftime('%d %b')} - {end_date.strftime('%d %b %Y')}  ({total_days} days)"
    logo = _pdf_logo()
    if logo:
        story.append(logo)
        story.append(Spacer(1,2))
    story.append(RP(f"{label.upper()} SUMMARY REPORT",bold=True,size=16,color=CR,align=TA_CENTER))
    story.append(Spacer(1,4))
    story.append(RP(f"Site: {site}   -   {dlabel}",size=10,color=CYL,align=TA_CENTER))
    story.append(Spacer(1,4))
    story.append(HRFlowable(width=W,thickness=1.5,color=CR,spaceAfter=8))
    cw=[W*0.04,W*0.26,W*0.08,W*0.10,W*0.10,W*0.12,W*0.12,W*0.11]
    for dk in DEPT_ORDER:
        d=ALL_EMPLOYEES[dk]; rows=summary[dk]
        if not rows: continue
        story.append(Spacer(1,4))
        story.append(_pdf_banner(d,W))
        data=[[RP("#",bold=True,size=8,color=CYLL,align=TA_CENTER),
               RP("Employee Name",bold=True,size=8,color=CYLL),
               RP("ID",bold=True,size=8,color=CYLL,align=TA_CENTER),
               RP("Present",bold=True,size=8,color=CYLL,align=TA_CENTER),
               RP("Absent",bold=True,size=8,color=CYLL,align=TA_CENTER),
               RP("Avg In",bold=True,size=8,color=CYLL,align=TA_CENTER),
               RP("Avg Out",bold=True,size=8,color=CYLL,align=TA_CENTER),
               RP("Total Hrs",bold=True,size=8,color=CYLL,align=TA_CENTER)]]
        style=[("BACKGROUND",(0,0),(-1,0),CRM),("BOX",(0,0),(-1,-1),0.5,CMG),
               ("INNERGRID",(0,0),(-1,-1),0.5,CMG),("VALIGN",(0,0),(-1,-1),"MIDDLE"),
               ("TOPPADDING",(0,0),(-1,-1),4),("BOTTOMPADDING",(0,0),(-1,-1),4),
               ("LEFTPADDING",(0,0),(-1,-1),5),("RIGHTPADDING",(0,0),(-1,-1),5)]
        for idx,rec in enumerate(rows):
            fill=CCR
            tc=CGT if rec["days_present"]==0 else colors.black
            if rec["days_present"]==0: fill=CGB
            data.append([RP(str(idx+1),size=8,color=tc,align=TA_CENTER),
                         RP(rec["name"],size=8,color=tc),
                         RP(rec["id"],size=8,color=tc,align=TA_CENTER),
                         RP(str(rec["days_present"]),bold=True,size=8,
                            color=CGN if rec["days_present"]>0 else CGT,align=TA_CENTER),
                         RP(str(rec["days_absent"]),bold=(rec["days_absent"]>0),size=8,
                            color=CRD if rec["days_absent"]>0 else CGT,align=TA_CENTER),
                         RP(rec["avg_in"],size=8,color=tc,align=TA_CENTER),
                         RP(rec["avg_out"],size=8,color=tc,align=TA_CENTER),
                         RP(rec["total_hrs"],size=8,color=tc,align=TA_CENTER)])
            ri=len(data)-1
            style.append(("BACKGROUND",(0,ri),(-1,ri),fill))
        story.append(_pdf_table(data,style,cw))
    story.append(Spacer(1,8))
    story.append(HRFlowable(width=W,thickness=1,color=CR,spaceAfter=4))
    story.append(RP(f"Generated: {datetime.today().strftime('%d %B %Y')}",size=8,color=CYL))
    doc.build(story); buf.seek(0); return buf.read()

# ══════════════════════════════════════════════════════════════
# ── EXCEL BUILDER ────────────────────────────────────────────
# ══════════════════════════════════════════════════════════════
XT=Side(style="thin",color="D4A017"); XK=Side(style="medium",color="8B0000")
XCB=Border(left=XT,right=XT,top=XT,bottom=XT)
XHB=Border(left=XK,right=XK,top=XK,bottom=XK)

def xf(color="2D2D2D",bold=False,size=10): return Font(name="Arial",bold=bold,size=size,color=color)
def xfill(h): return PatternFill("solid",fgColor=h)
def xal(h="center",v="center"): return Alignment(horizontal=h,vertical=v,wrap_text=False)

def _xl_title(ws,title,subtitle,ncols):
    end_col = chr(64+ncols)
    ws.row_dimensions[1].height = 78
    ws.merge_cells(f"A1:{end_col}1")

    logo_path = get_logo_path()
    if logo_path:
        logo = XLImage(logo_path)
        original_width = float(logo.width or 1)
        original_height = float(logo.height or 1)
        logo.width = 150
        logo.height = int(150 * original_height / original_width)
        logo.anchor = "C1" if ncols <= 6 else "D1"
        ws.add_image(logo)

    ws.row_dimensions[2].height=26
    ws.row_dimensions[3].height=16
    ws.merge_cells(f"A2:{end_col}2")
    ws.merge_cells(f"A3:{end_col}3")
    c=ws.cell(2,1,title); c.font=Font(name="Arial",bold=True,size=16,color="8B0000"); c.alignment=xal()
    c=ws.cell(3,1,subtitle); c.font=Font(name="Arial",size=11,color="B8860B"); c.alignment=xal()

def _xl_banner(ws,cur,label,shift,bg_hex,ncols):
    ws.row_dimensions[cur].height=6; cur+=1
    ws.row_dimensions[cur].height=20
    for col in range(1,ncols+1): cell=ws.cell(cur,col); cell.fill=xfill(bg_hex); cell.border=XHB
    ws.cell(cur,1,f"  {label}   -   {shift}")
    ws.cell(cur,1).font=Font(name="Arial",bold=True,size=11,color="FFF3CD")
    ws.cell(cur,1).alignment=xal("left")
    ws.merge_cells(start_row=cur,start_column=1,end_row=cur,end_column=ncols)
    return cur+1

def _xl_hdr(ws,cur,labels,aligns,ncols):
    ws.row_dimensions[cur].height=18
    for i,(lbl,al) in enumerate(zip(labels,aligns)):
        cell=ws.cell(cur,i+1,lbl); cell.font=Font(name="Arial",bold=True,size=10,color="FFF3CD")
        cell.fill=xfill("C0392B"); cell.alignment=xal(al); cell.border=XHB
    return cur+1

def build_xlsx_daily(rows_by_date,start_date,end_date,site="Buguruni"):
    wb=Workbook(); ws=wb.active; ws.title="Daily"
    dlabel=(start_date.strftime("%d %B %Y") if start_date==end_date
            else f"{start_date.strftime('%d %b')} - {end_date.strftime('%d %b %Y')}")
    _xl_title(ws,"DAILY SHIFT REPORT",f"Site: {site}   -   {dlabel}",6)
    ws.row_dimensions[4].height=13; ws.merge_cells("A4:F4")
    c=ws.cell(4,1,"Check-In Key:   Early (green)   On Time (white)   Late (red)   Absent (grey)")
    c.font=Font(name="Arial",size=8,color="555555"); c.alignment=xal("left")
    cur=5; total=0; day=start_date
    HDR=["#","Employee Name","ID","Check In","Check Out","Hours Worked"]
    AL=["center","left","center","center","center","center"]
    while day<=end_date:
        if (end_date-start_date).days>0:
            ws.row_dimensions[cur].height=16; ws.merge_cells(f"A{cur}:F{cur}")
            c=ws.cell(cur,1,day.strftime("%A, %d %B %Y"))
            c.font=Font(name="Arial",bold=True,size=11,color="8B0000"); c.alignment=xal("left"); cur+=1
        day_rows=rows_by_date.get(day.strftime("%Y-%m-%d"),[])
        total+=len(day_rows); dept_data=build_dept_data(day_rows)
        for dk in DEPT_ORDER:
            d=ALL_EMPLOYEES[dk]; pd=dept_data[dk]
            if not pd["present"] and not pd["absent"]: continue
            cur=_xl_banner(ws,cur,d["label"],d["shift"],d["header_hex"],6)
            cur=_xl_hdr(ws,cur,HDR,AL,6)
            for idx,rec in enumerate(pd["present"]):
                ws.row_dimensions[cur].height=16
                fh="FFFFFF"
                ci_bg,ci_tc=checkin_fill(dk,rec["check_in"])
                co=rec["check_out"] if rec["check_out"]!="-" else ""
                hrs=rec["hours"] if rec["hours"]!="-" else ""
                for i,(val,al,fhx,fc) in enumerate(zip(
                        [idx+1,rec["name"],rec["employee_id"],rec["check_in"],co,hrs],
                        AL,[fh,fh,fh,ci_bg,fh,fh],["2D2D2D","2D2D2D","2D2D2D",ci_tc,"2D2D2D","2D2D2D"])):
                    cell=ws.cell(cur,i+1,val)
                    cell.font=Font(name="Arial",size=10,color=fc,bold=(i==3))
                    cell.fill=xfill(fhx); cell.alignment=xal(al); cell.border=XCB
                cur+=1
            for j,nm in enumerate(pd["absent"]):
                ws.row_dimensions[cur].height=16
                configured_id = get_configured_employee_id(dk, nm)
                for i,(val,al) in enumerate(zip([len(pd["present"])+j+1,nm,configured_id,"-","-","-"],AL)):
                    cell=ws.cell(cur,i+1,val); cell.font=xf("999999")
                    cell.fill=xfill("EEEEEE"); cell.alignment=xal(al); cell.border=XCB
                cur+=1
        day+=timedelta(days=1)
    cur+=1; ws.merge_cells(f"A{cur}:F{cur}")
    c=ws.cell(cur,1,f"Total records: {total}   |   Generated: {datetime.today().strftime('%d %B %Y')}")
    c.font=Font(name="Arial",bold=True,size=9,color="8B0000"); c.alignment=xal("right")
    for col,w in zip(["A","B","C","D","E","F"],[5,30,10,12,12,18]):
        ws.column_dimensions[col].width=w
    buf=io.BytesIO(); wb.save(buf); buf.seek(0); return buf.read()

def build_xlsx_summary(summary,total_days,start_date,end_date,label,site="Buguruni"):
    wb=Workbook(); ws=wb.active; ws.title=f"{label} Summary"
    dlabel=f"{start_date.strftime('%d %b')} - {end_date.strftime('%d %b %Y')}  ({total_days} days)"
    _xl_title(ws,f"{label.upper()} SUMMARY REPORT",f"Site: {site}   -   {dlabel}",8)
    cur=5
    HDR=["#","Employee Name","ID","Days Present","Days Absent","Avg Check-In","Avg Check-Out","Total Hours"]
    AL=["center","left","center","center","center","center","center","center"]
    for dk in DEPT_ORDER:
        d=ALL_EMPLOYEES[dk]; rows=summary[dk]
        if not rows: continue
        cur=_xl_banner(ws,cur,d["label"],d["shift"],d["header_hex"],8)
        cur=_xl_hdr(ws,cur,HDR,AL,8)
        for idx,rec in enumerate(rows):
            ws.row_dimensions[cur].height=16
            fh="FFFFFF"
            if rec["days_present"]==0: fh="EEEEEE"
            gc="999999" if rec["days_present"]==0 else "2D2D2D"
            pc="276221" if rec["days_present"]>0 else "999999"
            ac="CC0000" if rec["days_absent"]>0 else "2D2D2D"
            for i,(val,al,fc) in enumerate(zip(
                    [idx+1,rec["name"],rec["id"],rec["days_present"],rec["days_absent"],
                     rec["avg_in"],rec["avg_out"],rec["total_hrs"]],
                    AL,[gc,gc,gc,pc,ac,gc,gc,gc])):
                cell=ws.cell(cur,i+1,val)
                cell.font=Font(name="Arial",size=10,color=fc,bold=(i in (3,4) and val not in (0,"-")))
                cell.fill=xfill(fh); cell.alignment=xal(al); cell.border=XCB
            cur+=1
    cur+=1; ws.merge_cells(f"A{cur}:H{cur}")
    c=ws.cell(cur,1,f"Generated: {datetime.today().strftime('%d %B %Y')}")
    c.font=Font(name="Arial",bold=True,size=9,color="8B0000"); c.alignment=xal("right")
    for col,w in zip(["A","B","C","D","E","F","G","H"],[5,28,10,13,13,14,14,14]):
        ws.column_dimensions[col].width=w
    buf=io.BytesIO(); wb.save(buf); buf.seek(0); return buf.read()

# ══════════════════════════════════════════════════════════════
# ── STREAMLIT UI ─────────────────────────────────────────────
# ══════════════════════════════════════════════════════════════

# ── Persist selected site across reruns ───────────────────────
site_names = list(SITES.keys())
if "selected_site" not in st.session_state:
    st.session_state.selected_site = site_names[0]

# Resolve the currently active site before drawing the header.
site_cfg = SITES[st.session_state.selected_site]
ALL_EMPLOYEES = site_cfg["employees"]
DEPT_ORDER = [
    department
    for department in site_cfg["dept_order"]
    if not (
        st.session_state.selected_site == "Puma Upanga"
        and department == "Manager"
    )
]
DEPT_COLORS = site_cfg["dept_colors"]
device_ip = site_cfg["device_ip"]
username = site_cfg["username"]
password = site_cfg["password"]

# Header — original simple structure with no Settings button.
st.markdown(f"# {site_cfg['label']} Shift Reports")

st.markdown(
    f"<p style='color:{MUTED};margin-top:-10px'>Daily attendance reports — "
    f"<b>{site_cfg['label']}</b> access control device.</p>",
    unsafe_allow_html=True,
)

# One row of site choices. Clicking the site name changes the active site;
# there are no separate Switch buttons.
selected_site = st.radio(
    "Select site",
    site_names,
    index=site_names.index(st.session_state.selected_site),
    horizontal=True,
    label_visibility="collapsed",
    key="site_selector_row",
)
if selected_site != st.session_state.selected_site:
    st.session_state.selected_site = selected_site
    try:
        st.rerun()
    except AttributeError:
        st.experimental_rerun()

st.caption(f"Active device: {site_cfg['device_ip']}")

st.markdown("---")

# ── Date range + quick presets ────────────────────────────────
st.markdown("### Date Range")
preset_col, _, dc1, dc2 = st.columns([3,0.3,2,2])
with preset_col:
    preset = st.selectbox("Quick select", ["Custom","Today","Yesterday","This week","Last week","This month","Last month"], index=1)

today = local_now().date()
if preset=="Today":
    default_start=default_end=today
elif preset=="Yesterday":
    default_start=default_end=today-timedelta(days=1)
elif preset=="This week":
    default_start=today-timedelta(days=today.weekday()); default_end=today
elif preset=="Last week":
    default_start=today-timedelta(days=today.weekday()+7)
    default_end=today-timedelta(days=today.weekday()+1)
elif preset=="This month":
    default_start=today.replace(day=1); default_end=today
elif preset=="Last month":
    first_this=today.replace(day=1)
    default_end=first_this-timedelta(days=1)
    default_start=default_end.replace(day=1)
else:
    default_start=default_end=today

with dc1:
    start_date = st.date_input("From", value=default_start, max_value=today)
with dc2:
    end_date = st.date_input("To", value=default_end, max_value=today)

if end_date < start_date:
    st.warning("End date is before start date — swapping."); start_date,end_date=end_date,start_date

num_days=(end_date-start_date).days+1
is_multi=num_days>1
tag=str(start_date) if not is_multi else f"{start_date}_to_{end_date}"

# auto-detect label
if num_days>=28:   period_label="Monthly"
elif num_days>=7:  period_label="Weekly"
else:              period_label="Daily"

st.markdown(f"<small style='color:{MUTED}'>{num_days} day(s) selected — <b>{period_label}</b> report</small>",
            unsafe_allow_html=True)
st.markdown("---")

# ── Generate ──────────────────────────────────────────────────
if st.button("View / Generate Reports"):
    progress=st.progress(0.0); status=st.empty()

    def upd(val,msg): progress.progress(val); status.info(msg)

    status.info("Connecting to device...")
    events,err=fetch_events(device_ip,username,password,start_date,end_date,upd)
    if err: st.error(f"Failed to fetch data: {err}"); st.stop()

    progress.progress(0.6); status.info("Processing records...")
    rows=parse_by_day(events,start_date,end_date)
    rows_by_date=defaultdict(list)
    for r in rows: rows_by_date[r["date"]].append(r)

    # ── Summary cards ─────────────────────────────────────────
    progress.progress(0.65)
    total_present=len(rows)
    all_members=sum(len(ALL_EMPLOYEES[dk]["members"]) for dk in DEPT_ORDER)
    st.markdown("### Summary")
    m1,m2,m3,m4,m5=st.columns(5)
    m1.metric("Days", num_days)
    m2.metric("Total Records", total_present)
    m3.metric("Staff Roster", all_members)
    m4.metric("Avg / Day", total_present//max(num_days,1))
    unique_staff=len(set(r["name"] for r in rows))
    m5.metric("Unique Staff", unique_staff)

    # ── Per-dept breakdown (last day or whole range) ───────────
    last_rows=rows_by_date.get(end_date.strftime("%Y-%m-%d"),[])
    dept_data_last=build_dept_data(last_rows)
    st.markdown(f"**Department breakdown — {end_date.strftime('%d %B %Y')}**")
    department_cards = ["<div class='department-grid'>"]
    for dk in DEPT_ORDER:
        pd_=dept_data_last[dk]
        department_cards.append(
            f"<div class='department-card' style='background:{DEPT_COLORS[dk]}'>"
            f"<div style='color:#FFF3CD;font-weight:bold;font-size:11px'>{ALL_EMPLOYEES[dk]['label']}</div>"
            f"<div style='color:#FFFFFF;font-size:22px;font-weight:bold'>{len(pd_['present'])}</div>"
            f"<div style='color:#F2E7D5;font-size:10px'>of {len(pd_['present'])+len(pd_['absent'])}</div>"
            f"</div>"
        )
    department_cards.append("</div>")
    st.markdown("".join(department_cards),unsafe_allow_html=True)

    # ── Build files ───────────────────────────────────────────
    summary,total_days_n=build_summary(rows_by_date,start_date,end_date)

    progress.progress(0.70); status.info("Building daily Word document...")
    site_label = site_cfg["label"].replace(" ","_")
    site_display = site_cfg["label"]
    docx_daily  = build_docx_daily(rows_by_date,start_date,end_date,site_display)

    progress.progress(0.78); status.info("Building daily PDF...")
    pdf_daily   = build_pdf_daily(rows_by_date,start_date,end_date,site_display)

    progress.progress(0.84); status.info("Building daily Excel...")
    xlsx_daily  = build_xlsx_daily(rows_by_date,start_date,end_date,site_display)

    progress.progress(0.88); status.info(f"Building {period_label} summary...")
    docx_summ   = build_docx_summary(summary,total_days_n,start_date,end_date,period_label,site_display)
    pdf_summ    = build_pdf_summary(summary,total_days_n,start_date,end_date,period_label,site_display)
    xlsx_summ   = build_xlsx_summary(summary,total_days_n,start_date,end_date,period_label,site_display)

    progress.progress(1.0); status.success("Reports ready!")

    # ── Viewer and downloads ──────────────────────────────────
    st.markdown("### Report Options")
    st.caption("Choose a button below to view the report on screen or download report files.")
    viewer_tab, downloads_tab = st.tabs(["👁  View Reports", "⬇  Download Reports"])

    with viewer_tab:
        current_local = local_now()
        if start_date == end_date == current_local.date():
            st.info(
                "Today's checkout is shown only when an employee scans at the allocated "
                "shift ending time or later. Earlier or repeated scans remain blank."
            )

        st.markdown(
            """
            <div style="display:flex;gap:12px;align-items:center;flex-wrap:wrap;margin:4px 0 14px 0;">
              <span style="font-weight:700;">Check-in key:</span>
              <span style="background:#C6EFCE;color:#276221;padding:4px 10px;border:1px solid #9BCB9F;border-radius:4px;">Early</span>
              <span style="background:#FFFFFF;color:#333333;padding:4px 10px;border:1px solid #CCCCCC;border-radius:4px;">On time</span>
              <span style="background:#FFCCCC;color:#CC0000;padding:4px 10px;border:1px solid #E6A0A0;border-radius:4px;">Late</span>
              <span style="background:#EEEEEE;color:#777777;padding:4px 10px;border:1px solid #CCCCCC;border-radius:4px;">Absent</span>
            </div>
            """,
            unsafe_allow_html=True,
        )

        daily_view_tab, summary_view_tab = st.tabs([
            f"Daily Attendance ({num_days} day{'s' if num_days > 1 else ''})",
            f"{period_label} Summary",
        ])

        with daily_view_tab:
            view_day = start_date
            while view_day <= end_date:
                day_key = view_day.strftime("%Y-%m-%d")
                st.markdown(f"#### {view_day.strftime('%A, %d %B %Y')}")
                day_data = build_dept_data(rows_by_date.get(day_key, []))

                for dk in DEPT_ORDER:
                    department = ALL_EMPLOYEES[dk]
                    attendance = day_data[dk]
                    st.markdown(
                        f"**{department['label']}** — {department['shift']} "
                        f"({len(attendance['present'])} present, {len(attendance['absent'])} absent)"
                    )

                    viewer_rows = []
                    for rec in attendance["present"]:
                        viewer_rows.append({
                            "Status": "Present",
                            "Employee Name": rec["name"],
                            "ID": rec["employee_id"],
                            "Check In": rec["check_in"],
                            "Check Out": "" if rec["check_out"] == "-" else rec["check_out"],
                            "Hours Worked": "" if rec["hours"] == "-" else rec["hours"],
                        })
                    for absent_name in attendance["absent"]:
                        viewer_rows.append({
                            "Status": "Absent",
                            "Employee Name": absent_name,
                            "ID": get_configured_employee_id(dk, absent_name),
                            "Check In": "",
                            "Check Out": "",
                            "Hours Worked": "",
                        })

                    if viewer_rows:
                        viewer_columns = [
                            "Status", "Employee Name", "ID",
                            "Check In", "Check Out", "Hours Worked",
                        ]
                        viewer_df = pd.DataFrame(viewer_rows, columns=viewer_columns)
                        styled_viewer = style_viewer_dataframe(viewer_df, dk)
                        st.dataframe(
                            styled_viewer,
                            hide_index=True,
                            use_container_width=True,
                            column_order=viewer_columns,
                        )
                view_day += timedelta(days=1)

        with summary_view_tab:
            st.caption(
                "One row per employee showing attendance totals and average times "
                "for the selected period."
            )
            for dk in DEPT_ORDER:
                department = ALL_EMPLOYEES[dk]
                st.markdown(f"**{department['label']}** — {department['shift']}")
                summary_rows = []
                for rec in summary[dk]:
                    summary_rows.append({
                        "Employee Name": rec["name"],
                        "ID": rec["id"],
                        "Days Present": rec["days_present"],
                        "Days Absent": rec["days_absent"],
                        "Average Check In": "" if rec["avg_in"] == "-" else rec["avg_in"],
                        "Average Check Out": "" if rec["avg_out"] == "-" else rec["avg_out"],
                        "Total Hours": "" if rec["total_hrs"] == "-" else rec["total_hrs"],
                    })
                st.dataframe(summary_rows, hide_index=True, use_container_width=True)

    with downloads_tab:
        daily_download_tab, summary_download_tab = st.tabs([
            f"Daily ({num_days} day{'s' if num_days > 1 else ''})",
            f"{period_label} Summary",
        ])

        with daily_download_tab:
            d1,d2,d3=st.columns(3)
            with d1:
                st.download_button("⬇ Download Word (.docx)",data=docx_daily,
                    file_name=f"{site_label}_Daily_{tag}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            with d2:
                st.download_button("⬇ Download PDF",data=pdf_daily,
                    file_name=f"{site_label}_Daily_{tag}.pdf",mime="application/pdf")
            with d3:
                st.download_button("⬇ Download Excel (.xlsx)",data=xlsx_daily,
                    file_name=f"{site_label}_Daily_{tag}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

        with summary_download_tab:
            st.markdown(f"<small style='color:{MUTED}'>One row per employee — days present/absent, average check-in/out, total hours worked across the selected period.</small>",
                        unsafe_allow_html=True)
            d1,d2,d3=st.columns(3)
            with d1:
                st.download_button("⬇ Download Word (.docx)",data=docx_summ,
                    file_name=f"{site_label}_{period_label}Summary_{tag}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            with d2:
                st.download_button("⬇ Download PDF",data=pdf_summ,
                    file_name=f"{site_label}_{period_label}Summary_{tag}.pdf",mime="application/pdf")
            with d3:
                st.download_button("⬇ Download Excel (.xlsx)",data=xlsx_summ,
                    file_name=f"{site_label}_{period_label}Summary_{tag}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

st.markdown("---")
st.markdown(f"<small style='color:{MUTED}'>{site_cfg['label']} Shift Report System</small>",unsafe_allow_html=True)
